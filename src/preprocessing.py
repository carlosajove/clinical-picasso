from dataclasses import dataclass
from pathlib import Path
import hashlib
import io


@dataclass
class DocumentRecord:
    filename: str
    size_bytes: int
    sha256: str
    raw_bytes: bytes
    is_duplicate: bool = False
    content: str | bytes | None = None


class Preprocessing:
    def __init__(self, data_dir: str | Path):
        self.data_dir = Path(data_dir)
        self.documents: list[DocumentRecord] = []
        self.kept: list[DocumentRecord] = []
        self.removed: list[DocumentRecord] = []

    def load_all(self) -> "Preprocessing":
        """Read every file in data_dir, compute SHA-256 hash."""
        for p in sorted(self.data_dir.iterdir()):
            if p.is_dir() or p.name.startswith(".") or p.name.startswith("~$"):
                continue
            raw = p.read_bytes()
            self.documents.append(
                DocumentRecord(
                    filename=p.name,
                    size_bytes=len(raw),
                    sha256=hashlib.sha256(raw).hexdigest(),
                    raw_bytes=raw,
                )
            )
        return self

    def deduplicate(self) -> "Preprocessing":
        """Group by SHA-256. Keep first occurrence, mark the rest as duplicates."""
        seen: dict[str, DocumentRecord] = {}
        for doc in self.documents:
            if doc.sha256 not in seen:
                seen[doc.sha256] = doc
                self.kept.append(doc)
            else:
                doc.is_duplicate = True
                self.removed.append(doc)
        return self

    def deduplicate_near(self, threshold: float = 0.9) -> "Preprocessing":
        """Second pass: mark near-duplicates where both cosine similarity
        and size ratio exceed the threshold."""
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        import numpy as np

        texts = [self._extract_text(doc) for doc in self.kept]
        sizes = np.array([len(t) for t in texts], dtype=float)

        tfidf = TfidfVectorizer(max_features=5000).fit_transform(texts)
        cos_sim = cosine_similarity(tfidf)

        max_sizes = np.maximum(sizes[:, None], sizes[None, :])
        max_sizes[max_sizes == 0] = 1
        size_ratio = np.minimum(sizes[:, None], sizes[None, :]) / max_sizes

        to_remove = set()
        for i in range(len(self.kept)):
            if i in to_remove:
                continue
            for j in range(i + 1, len(self.kept)):
                if j in to_remove:
                    continue
                if cos_sim[i, j] > threshold and size_ratio[i, j] > threshold:
                    to_remove.add(j)

        new_kept = []
        for i, doc in enumerate(self.kept):
            if i in to_remove:
                doc.is_duplicate = True
                self.removed.append(doc)
            else:
                new_kept.append(doc)
        self.kept = new_kept
        return self

    def _extract_text(self, doc: "DocumentRecord") -> str:
        ext = Path(doc.filename).suffix.lower()
        try:
            if ext in (".txt", ".md", ".csv"):
                return doc.raw_bytes.decode("utf-8", errors="replace")
            if ext == ".html":
                from bs4 import BeautifulSoup
                return BeautifulSoup(doc.raw_bytes, "html.parser").get_text(separator="\n")
            if ext == ".pdf":
                import fitz
                pdf = fitz.open(stream=doc.raw_bytes, filetype="pdf")
                return "\n".join(page.get_text() for page in pdf)
            if ext == ".docx":
                import docx
                d = docx.Document(io.BytesIO(doc.raw_bytes))
                return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:
            print(f"Warning: text extraction failed for {doc.filename}: {e}")
            return ""
        return doc.raw_bytes.decode("utf-8", errors="replace")

    def extract_content(self) -> "Preprocessing":
        """Populate content: str for text formats, bytes (PDF) for pdf/docx."""
        for doc in self.kept:
            ext = Path(doc.filename).suffix.lower()
            if ext == ".pdf":
                doc.content = doc.raw_bytes
            elif ext == ".docx":
                doc.content = self._docx_to_pdf(doc)
            else:
                doc.content = self._extract_text(doc)
        return self

    def _docx_to_pdf(self, doc: "DocumentRecord") -> bytes:
        import docx
        from fpdf import FPDF

        d = docx.Document(io.BytesIO(doc.raw_bytes))
        pdf = FPDF()
        pdf.set_margin(10)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=10)
        font_path = "/System/Library/Fonts/Helvetica.ttc"
        if Path(font_path).exists():
            pdf.add_font("uni", "", font_path, uni=True)
            pdf.set_font("uni", size=9)
        else:
            pdf.set_font("Helvetica", size=9)
        w = pdf.w - pdf.l_margin - pdf.r_margin
        for para in d.paragraphs:
            text = para.text.strip()
            if text:
                pdf.multi_cell(w, 4, text)
        for table in d.tables:
            for row in table.rows:
                line = " | ".join(cell.text for cell in row.cells)
                if line.strip():
                    pdf.multi_cell(w, 4, line)
        return bytes(pdf.output())

    def summary(self) -> dict:
        by_ext = {}
        for doc in self.documents:
            ext = Path(doc.filename).suffix.lower()
            by_ext[ext] = by_ext.get(ext, 0) + 1
        return {
            "total_files": len(self.documents),
            "unique": len(self.kept),
            "duplicates": len(self.removed),
            "by_extension": by_ext,
            "duplicates_removed": [d.filename for d in self.removed],
        }
